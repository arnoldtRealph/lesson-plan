import streamlit as st
<<<<<<< HEAD
import io
import docx
from datetime import date
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
=======
import os
import io
import docx
from datetime import date
import streamlit_analytics

>>>>>>> 3316b4cdb103047103c520996e340b72d4467e3b


# Set page title and icon
st.set_page_config(page_title="Lesson Plan Creator", page_icon=":books:")

<<<<<<< HEAD
# Create input fields
st.title("SAUL DAMON HIGH SCHOOL")
st.subheader("LESSON PLANNER")
st.subheader("Please fill in the following fields to create a new lesson plan:")

=======

# Create input fields
st.title("SAUL DAMON HIGH SCHOOL")
st.subheader("LESSON PLANNER")
st.write("Please fill in the following fields to create a new lesson plan:")
st.write("")
>>>>>>> 3316b4cdb103047103c520996e340b72d4467e3b
subject = st.text_input("SUBJECT")
lesson_title = st.text_input("LESSON TITLE")
grade = st.selectbox("GRADE", ["9", "10", "11", "12"])
st.subheader("CHOOSE YOUR START AND END DATE")
start_date = st.date_input("FROM", value=date.today())
end_date = st.date_input("TO", value=date.today())
lesson_objective = st.text_area("LESSON OBJECTIVES")
lesson_introduction = st.text_area("INTRODUCTION")
lesson_activities = st.text_area("LESSON ACTIVITIES")
materials_needed = st.text_area("MATERIAL NEEDED")
homework = st.text_area("HOMEWORK ACTIVITIES")
notes = st.text_area("NOTES")
<<<<<<< HEAD

st.subheader("TEACHER INFORMATION")
teacher_name = st.text_input("INITIALS")
teacher_surname = st.text_input("SURNAME")

=======
st.subheader("TEACHER INFORMATION")
teacher_name = st.text_input("INITIALS")
teacher_surname = st.text_input("SURNAME")
st.write("")
>>>>>>> 3316b4cdb103047103c520996e340b72d4467e3b
st.write("Created by Mr. A.R Visagie @ Saul Damon High School")

# Create save button
if st.button("Create Lesson Plan"):
    # Create a new Word document
    document = docx.Document()
<<<<<<< HEAD

    # Add subject as a bold, navy-colored heading
    title = document.add_heading(level=0)
    run = title.add_run(subject.upper())  # Add text to the heading
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 128)  # Navy color

    # Add lesson details in a table
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Pt(180)
    table.columns[1].width = Pt(300)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "LESSON TITLE:"
    hdr_cells[1].text = lesson_title
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "GRADE:"
    hdr_cells[1].text = grade
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "FROM:"
    hdr_cells[1].text = str(start_date)
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "TO:"
    hdr_cells[1].text = str(end_date)

    # Add section headings and their contents in formal format
    document.add_paragraph()
    document.add_heading("LESSON OBJECTIVES", level=1)
    document.add_paragraph(lesson_objective)
    
    document.add_heading("INTRODUCTION", level=1)
    document.add_paragraph(lesson_introduction)
    
    document.add_heading("LESSON ACTIVITIES", level=1)
    document.add_paragraph(lesson_activities)
    
    document.add_heading("MATERIAL NEEDED", level=1)
    document.add_paragraph(materials_needed)
    
    document.add_heading("HOMEWORK ACTIVITIES", level=1)
    document.add_paragraph(homework)
    
    document.add_heading("NOTES", level=1)
    document.add_paragraph(notes)
    
    # Add teacher information in a table for clean layout
    document.add_paragraph()
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Pt(180)
    table.columns[1].width = Pt(300)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "INITIALS:"
    hdr_cells[1].text = teacher_name
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "SURNAME:"
    hdr_cells[1].text = teacher_surname
    hdr_cells = table.add_row().cells
    hdr_cells[0].text = "SIGNATURE:"
    hdr_cells[1].text = ""
=======
    # Add input values to the document
    document.add_heading(subject.upper(), level=0)
    document.add_paragraph("")
    document.add_paragraph("LESSON TITLE: " + lesson_title)
    document.add_paragraph("GRADE: " + grade)
    document.add_paragraph("FROM: " + str(start_date))
    document.add_paragraph("TO: " + str(end_date))
    document.add_paragraph("")
    document.add_heading("LESSON OBJECTIVES", level=1)
    document.add_paragraph(lesson_objective)
    document.add_heading("INTRODUCTION", level=1)
    document.add_paragraph(lesson_introduction)
    document.add_heading("LESSON ACTIVITIES", level=1)
    document.add_paragraph(lesson_activities)
    document.add_heading("MATERIAL NEEDED", level=1)
    document.add_paragraph(materials_needed)
    document.add_heading("HOMEWORK ACTIVITIES", level=1)
    document.add_paragraph(homework)
    document.add_heading("NOTES", level=1)
    document.add_paragraph(notes)
    document.add_paragraph("")
    document.add_paragraph("INITIALS: " + teacher_name)
    document.add_paragraph("SURNAME: " + teacher_surname)
    document.add_paragraph("SIGNATURE: ")
>>>>>>> 3316b4cdb103047103c520996e340b72d4467e3b

    # Save document to BytesIO object
    with io.BytesIO() as output:
        document.save(output)
        output.seek(0)
<<<<<<< HEAD
        
=======
>>>>>>> 3316b4cdb103047103c520996e340b72d4467e3b
        # Create download button
        st.download_button(
            label="Download Lesson Plan",
            data=output,
            file_name="lesson_plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("Your lesson plan has been created. Click the download button to save the file.")
<<<<<<< HEAD
=======


streamlit_analytics.stop_tracking()
>>>>>>> 3316b4cdb103047103c520996e340b72d4467e3b
